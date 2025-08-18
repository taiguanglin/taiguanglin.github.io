#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
问答录目录转Word文档程序
将wenda2-toc-2025-08-18v2.txt文件转换为Word文档格式的目录
"""

import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import os

class WendaTocConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
    
    def setup_document(self):
        """设置文档基本格式"""
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # 添加标题
        title = self.doc.add_heading('问答录目录', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        self.doc.add_paragraph()
    
    def parse_line(self, line):
        """解析目录行，返回层级、标题和条目数"""
        original_line = line
        line = line.rstrip()
        if not line:
            return None
            
        # 匹配模式：层级标识 + 标题 + (条目数)
        pattern = r'^(\s*)(.*?)\s+\((\d+)\)$'
        match = re.match(pattern, line)
        
        if not match:
            return None
            
        indent_spaces = match.group(1)
        title = match.group(2).strip()
        item_count = match.group(3)  # 改為條目數，不是頁碼
        
        # 基于缩进和内容确定层级
        indent_count = len(indent_spaces)
        
        # 主要章节判断 (I. II. III. 等)
        if re.match(r'^[IVX]+\.\s+', title):
            level = 0
        # 二级章节判断 (I.I. I.II. 等)
        elif re.match(r'^[IVX]+\.[IVX]+\.\s+', title):
            level = 1
        else:
            # 基於縮排確定層級（這是最關鍵的判斷標準）
            if indent_count == 0:
                level = 0
            elif indent_count <= 2:
                level = 1  
            elif indent_count <= 4:
                level = 2  # 所有4個空格縮排的項目（包括數字開頭和箭頭項目）
            elif indent_count <= 6:
                level = 3
            elif indent_count <= 8:
                level = 4
            else:
                level = 5
        

        return {
            'level': level,
            'title': title,
            'item_count': item_count,  # 改為條目數
            'indent': indent_count,
            'original_line': original_line.strip()
        }
    
    def add_toc_entry(self, entry):
        """添加目录条目到Word文档"""
        if not entry:
            return
            
        level = entry['level']
        title = entry['title']
        item_count = entry['item_count']  # 條目數
        
        # 创建段落
        p = self.doc.add_paragraph()
        
        # 根据层级设置缩进和格式
        if level == 0:  # 主要章节 (I, II, III...)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(f"{title}")
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = '微软雅黑'
            
        elif level == 1:  # 二级章节 (I.I, I.II...) - 與主章節對齊
            p.paragraph_format.left_indent = Cm(0)  # 修正：不縮排
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"{title}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = '微软雅黑'
            
        elif level == 2:  # 三级章节 - 適度縮排
            p.paragraph_format.left_indent = Cm(0.8)  # 輕微縮排
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{title}")
            run.font.size = Pt(12)
            run.font.name = '宋体'
            
        elif level == 3:  # 四级章节
            p.paragraph_format.left_indent = Cm(1.6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{title}")
            run.font.size = Pt(11)
            run.font.name = '宋体'
            
        else:  # 更深层级
            indent_cm = 2.4 + (level - 4) * 0.8
            p.paragraph_format.left_indent = Cm(indent_cm)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"{title}")
            run.font.size = Pt(10)
            run.font.name = '宋体'
        
        # 添加條目數（簡潔格式，無點點點）
        count_run = p.add_run(f" ({item_count}條)")
        
        # 根據層級設定條目數的字體大小
        if level <= 1:
            count_run.font.size = Pt(12)
        else:
            count_run.font.size = Pt(10)
            
        count_run.font.name = '宋体'
        count_run.italic = True
    
    def convert_file(self, input_file, output_file):
        """转换目录文件为Word文档"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            print(f"正在处理 {len(lines)} 行目录数据...")
            
            total_entries = 0
            total_item_count = 0
            
            for i, line in enumerate(lines, 1):
                entry = self.parse_line(line)
                if entry:
                    self.add_toc_entry(entry)
                    total_entries += 1
                    total_item_count += int(entry['item_count'])  # 累計條目數
                    
                    if i % 50 == 0:
                        print(f"已处理 {i}/{len(lines)} 行，有效条目 {total_entries} 个")
            
            # 添加统计信息
            self.add_summary_info(total_entries, total_item_count)
            
            # 保存文档
            self.doc.save(output_file)
            print(f"转换完成！输出文件：{output_file}")
            print(f"总共处理了 {total_entries} 个目录分类，包含 {total_item_count} 个条目")
            
        except FileNotFoundError:
            print(f"错误：找不到输入文件 {input_file}")
        except Exception as e:
            print(f"转换过程中发生错误：{str(e)}")
    
    def add_summary_info(self, total_entries, total_item_count):
        """添加统计信息"""
        from datetime import datetime
        
        # 添加分页符
        self.doc.add_page_break()
        
        # 添加统计标题
        summary = self.doc.add_heading('目录统计信息', level=1)
        summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加统计内容
        stats_p = self.doc.add_paragraph()
        stats_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        stats_text = f"""
📊 目录统计
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

目录分类數：{total_entries:,} 個
總條目數：{total_item_count:,} 條
生成時間：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}

說明：括號內數字代表該分類包含的條目數量，不是頁碼

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        """
        
        run = stats_p.add_run(stats_text.strip())
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'

def main():
    """主函数"""
    import datetime
    
    # 输入和输出文件路径
    input_file = "wenda2-toc-2025-08-18v2.txt"
    output_file = "问答录目录.docx"
    
    print("问答录目录转Word文档程序")
    print("=" * 40)
    
    # 检查输入文件是否存在
    if not os.path.exists(input_file):
        print(f"错误：找不到输入文件 {input_file}")
        return
    
    # 创建转换器并执行转换
    converter = WendaTocConverter()
    converter.convert_file(input_file, output_file)
    
    print("\n转换完成！")
    print(f"输出文件：{output_file}")

if __name__ == "__main__":
    main()
